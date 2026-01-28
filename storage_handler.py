import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
from datetime import datetime
from bs4 import BeautifulSoup

def save_to_excel(article, keyword, metadata=None):
    """
    يحفظ بيانات المقال والمنافسين في ملف Excel.
    """
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    
    data = {
        "الكلمة المفتاحية": [keyword],
        "العنوان": [article.get("title", keyword)],
        "عدد الكلمات": [article.get("word_count", 0)],
        "التاريخ": [timestamp],
        "الحالة": ["تم التوليد"],
        "الدومين المستهدف": [metadata.get("target_domain", "") if metadata else ""],
        "اللغة": [metadata.get("language", "العربية") if metadata else "العربية"]
    }
    
    df = pd.DataFrame(data)
    file_path = f"/tmp/{keyword.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.xlsx"
    
    try:
        df.to_excel(file_path, index=False, sheet_name="المقالات")
        return file_path
    except Exception as e:
        print(f"خطأ في حفظ Excel: {e}")
        return None

def save_to_docx(article, keyword, metadata=None):
    """
    يحفظ محتوى المقال في ملف Word مع تنسيق احترافي.
    """
    doc = Document()
    
    # إضافة العنوان الرئيسي
    title = doc.add_heading(article.get("title", keyword), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # إضافة معلومات المقال
    if metadata:
        info_para = doc.add_paragraph()
        info_para.add_run(f"الكلمة المفتاحية: ").bold = True
        info_para.add_run(metadata.get("main_keyword", ""))
        
        info_para.add_run("\n")
        info_para.add_run(f"تاريخ الإنشاء: ").bold = True
        info_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
    
    # تحويل HTML إلى نص منسق
    html_content = article.get("html", "")
    if html_content:
        try:
            soup = BeautifulSoup(html_content, "html.parser")
            
            for element in soup.find_all(['h1', 'h2', 'h3', 'p', 'li', 'a']):
                if element.name == 'h1':
                    heading = doc.add_heading(element.get_text(), level=1)
                elif element.name == 'h2':
                    heading = doc.add_heading(element.get_text(), level=2)
                elif element.name == 'h3':
                    heading = doc.add_heading(element.get_text(), level=3)
                elif element.name == 'p':
                    para = doc.add_paragraph(element.get_text())
                    para.paragraph_format.line_spacing = 1.5
                elif element.name == 'li':
                    doc.add_paragraph(element.get_text(), style='List Bullet')
                elif element.name == 'a':
                    para = doc.add_paragraph()
                    run = para.add_run(element.get_text())
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                    
        except Exception as e:
            print(f"خطأ في تحويل HTML: {e}")
            doc.add_paragraph(html_content)
    
    # إضافة Meta Description في النهاية
    if "meta_description" in article:
        doc.add_heading("وصف الميتا", level=2)
        doc.add_paragraph(article["meta_description"])
    
    file_path = f"/tmp/{keyword.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    
    try:
        doc.save(file_path)
        return file_path
    except Exception as e:
        print(f"خطأ في حفظ Docx: {e}")
        return None

def save_to_html(article, keyword):
    """
    يحفظ المقال كملف HTML مستقل.
    """
    html_template = f"""
    <!DOCTYPE html>
    <html dir="rtl" lang="ar">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>{article.get('title', keyword)}</title>
        <meta name="description" content="{article.get('meta_description', '')}">
        <style>
            body {{
                font-family: 'Cairo', Arial, sans-serif;
                line-height: 1.8;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                background-color: #f5f5f5;
                color: #333;
            }}
            h1, h2, h3 {{
                color: #2c3e50;
                margin-top: 20px;
                margin-bottom: 10px;
            }}
            a {{
                color: #3498db;
                text-decoration: none;
            }}
            a:hover {{
                text-decoration: underline;
            }}
            .meta-info {{
                background-color: #ecf0f1;
                padding: 10px;
                border-right: 4px solid #3498db;
                margin: 20px 0;
            }}
        </style>
    </head>
    <body>
        <h1>{article.get('title', keyword)}</h1>
        <div class="meta-info">
            <p><strong>عدد الكلمات:</strong> {article.get('word_count', 0)}</p>
            <p><strong>تاريخ الإنشاء:</strong> {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
        </div>
        {article.get('html', '')}
    </body>
    </html>
    """
    
    file_path = f"/tmp/{keyword.replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(html_template)
        return file_path
    except Exception as e:
        print(f"خطأ في حفظ HTML: {e}")
        return None

def upload_to_google_drive(file_path, folder_id=None):
    """
    يرفع الملف إلى Google Drive.
    يتطلب مفتاح API وتفويض المستخدم.
    """
    try:
        from google.colab import auth
        from googleapiclient.discovery import build
        from googleapiclient.http import MediaFileUpload
        
        auth.authenticate_user()
        drive_service = build('drive', 'v3')
        
        file_metadata = {'name': os.path.basename(file_path)}
        if folder_id:
            file_metadata['parents'] = [folder_id]
        
        media = MediaFileUpload(file_path, mimetype='application/octet-stream')
        file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
        
        return f"تم الرفع بنجاح! معرف الملف: {file.get('id')}"
    except Exception as e:
        return f"خطأ في الرفع إلى Google Drive: {str(e)}"
